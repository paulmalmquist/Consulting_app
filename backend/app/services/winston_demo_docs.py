from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from xml.sax.saxutils import escape
import zipfile

try:
    from docx import Document
except ModuleNotFoundError:
    Document = None

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
except ModuleNotFoundError:
    LETTER = None
    canvas = None


@dataclass(frozen=True)
class GeneratedDemoDocument:
    path: Path
    doc_type: str
    author: str
    linked_entities: list[dict[str, str]]


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _app_root() -> Path:
    return Path(__file__).resolve().parent.parent


def fixture_path() -> Path:
    return _app_root() / "fixtures" / "winston_demo" / "meridian_demo_seed.json"


def load_demo_fixture() -> dict:
    return json.loads(fixture_path().read_text(encoding="utf-8"))


def demo_docs_dir() -> Path:
    return _repo_root() / "demo_docs"


def _money(value: int | float) -> str:
    return f"${value:,.0f}"


def _pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def _write_pdf(path: Path, lines: list[str]) -> None:
    if canvas is None or LETTER is None:
        _write_basic_pdf(path, lines)
        return

    pdf = canvas.Canvas(str(path), pagesize=LETTER)
    _width, height = LETTER
    y = height - 54
    for line in lines:
        for chunk in _wrap(line, 92):
            pdf.drawString(54, y, chunk)
            y -= 14
            if y < 54:
                pdf.showPage()
                y = height - 54
    pdf.save()


def _write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    if Document is None:
        _write_basic_docx(path, title, paragraphs)
        return

    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def _write_basic_pdf(path: Path, lines: list[str]) -> None:
    wrapped_lines: list[str] = []
    for line in lines:
        wrapped_lines.extend(_wrap(line, 92))

    commands = ["BT", "/F1 12 Tf", "54 738 Td"]
    first = True
    for line in wrapped_lines:
        if not first:
            commands.append("0 -14 Td")
        first = False
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        commands.append(f"({escaped}) Tj")
    commands.append("ET")

    stream = "\n".join(commands).encode("utf-8")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Count 1 /Kids [3 0 R] >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n",
        f"4 0 obj << /Length {len(stream)} >> stream\n".encode("utf-8") + stream + b"\nendstream\nendobj\n",
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
    ]

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)
    xref_start = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("utf-8"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("utf-8"))
    pdf.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF\n".encode("utf-8")
    )
    path.write_bytes(bytes(pdf))


def _write_basic_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    body = []
    for text in [title, *paragraphs]:
        body.append(
            "<w:p><w:r><w:t xml:space=\"preserve\">"
            + escape(text)
            + "</w:t></w:r></w:p>"
        )
    document_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        "<w:body>"
        + "".join(body)
        + "<w:sectPr><w:pgSz w:w=\"12240\" w:h=\"15840\" /><w:pgMar "
        "w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" "
        "w:header=\"720\" w:footer=\"720\" w:gutter=\"0\" /></w:sectPr>"
        "</w:body></w:document>"
    )
    content_types = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
        "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
        "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
        "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
        "<Override PartName=\"/word/document.xml\" "
        "ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
        "</Types>"
    )
    package_rels = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
        "<Relationship Id=\"rId1\" "
        "Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" "
        "Target=\"word/document.xml\"/>"
        "</Relationships>"
    )
    document_rels = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\"></Relationships>"
    )

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", package_rels)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", document_rels)


def _write_text(path: Path, blocks: list[str]) -> None:
    path.write_text("\n\n".join(blocks) + "\n", encoding="utf-8")


def _write_csv(path: Path, headers: list[str], rows: list[list[str]]) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(headers)
        writer.writerows(rows)


def _wrap(text: str, width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if len(candidate) <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def generate_demo_docs(output_dir: Path | None = None) -> list[GeneratedDemoDocument]:
    fixture = load_demo_fixture()
    env = fixture["environment"]
    fund = fixture["fund"]
    assets = fixture["assets"]
    target_dir = output_dir or demo_docs_dir()
    target_dir.mkdir(parents=True, exist_ok=True)

    docs: list[GeneratedDemoDocument] = []

    def add(path: Path, doc_type: str, author: str, linked_entities: list[dict[str, str]]) -> None:
        docs.append(
            GeneratedDemoDocument(
                path=path,
                doc_type=doc_type,
                author=author,
                linked_entities=linked_entities,
            )
        )

    # Investor letter
    investor_letter = target_dir / "01-investor-letter-q1-2026.pdf"
    _write_pdf(
        investor_letter,
        [
            f"{env['client_name']} | Q1 2026 Investor Letter",
            f"{fund['name']} reported Q1 2026 total NOI of {_money(fund['total_noi'])}, "
            f"portfolio NAV of {_money(fund['portfolio_nav'])}, TVPI of {fund['tvpi']:.2f}x, "
            f"gross IRR of {_pct(fund['gross_irr'])}, and net IRR of {_pct(fund['net_irr'])}.",
            "Management notes that valuation remains supported by disciplined cap rate assumptions and "
            "stable property-level coverage metrics across multifamily, senior housing, student housing, "
            "medical office, and industrial holdings.",
        ],
    )
    add(investor_letter, "Investor Letter", "Investor Relations", [{"type": "fund", "id": fund["fund_id"]}])

    # Valuation policy
    valuation_policy = target_dir / "02-valuation-policy-memo.pdf"
    _write_pdf(
        valuation_policy,
        [
            "Valuation Policy Memo - Q1 2026",
            "The committee applies direct capitalization as the primary valuation method for the Winston "
            "institutional demo. NOI is annualized and triangulated against transaction comps and sector risk.",
            "For the seeded downside case, a +75 bps exit cap rate adjustment reduces portfolio NAV by "
            f"{_money(abs(fixture['downside']['portfolio_nav_delta']))} and compresses TVPI by "
            f"{fixture['downside']['tvpi_delta']:.2f} turns.",
        ],
    )
    add(valuation_policy, "Valuation Policy Memo", "Valuation Committee", [{"type": "fund", "id": fund["fund_id"]}])

    # Asset-specific docs
    for index, asset in enumerate(assets, start=1):
        linked = [
            {"type": "fund", "id": fund["fund_id"]},
            {"type": "asset", "id": asset["asset_id"]},
        ]
        memo_path = target_dir / f"{2 + index:02d}-underwriting-{index}-{asset['name'].lower().replace(' ', '-')}.docx"
        _write_docx(
            memo_path,
            f"{asset['name']} Underwriting Memo",
            [
                f"{asset['name']} is categorized as {asset['property_type']} in {asset['market']}.",
                f"Q1 2026 NOI is {_money(asset['noi'])}, stabilized value is {_money(asset['asset_value'])}, "
                f"and debt outstanding is {_money(asset['debt_balance'])}.",
                f"Current DSCR is {asset['dscr']:.2f}x and WALT is {asset['walt']:.1f} years.",
                "The underwriting conclusion supports the current valuation so long as NOI and occupancy "
                "remain within the underwritten range.",
            ],
        )
        add(memo_path, "Asset Underwriting Memo", "Asset Management", linked)

        transcript_path = target_dir / f"{7 + index:02d}-operating-call-{index}-{asset['name'].lower().replace(' ', '-')}.txt"
        _write_text(
            transcript_path,
            [
                f"Operating Call Transcript - {asset['name']} - Q1 2026",
                f"Asset manager: {asset['name']} delivered NOI of {_money(asset['noi'])}. "
                f"Coverage held at {asset['dscr']:.2f}x and valuation was marked at {_money(asset['asset_value'])}.",
                "Chief investment officer: Maintain focus on expense control, leasing velocity, and debt service "
                "coverage through the next quarter close.",
            ],
        )
        add(transcript_path, "Operating Call Transcript", "Operations Desk", linked)

        utilities_path = target_dir / f"{17 + index:02d}-utilities-{index}-{asset['name'].lower().replace(' ', '-')}.csv"
        _write_csv(
            utilities_path,
            ["asset_name", "quarter", "electricity", "water", "gas", "notes"],
            [
                [asset["name"], fund["quarter"], "84250", "21180", "10840", "Utilities aligned to Q1 NOI bridge"],
            ],
        )
        add(utilities_path, "Utility Expense Summary", "Property Accounting", linked)

        esg_path = target_dir / f"{22 + index:02d}-esg-{index}-{asset['name'].lower().replace(' ', '-')}.txt"
        _write_text(
            esg_path,
            [
                f"ESG Snapshot - {asset['name']}",
                f"{asset['name']} reported Q1 2026 operating performance with NOI of {_money(asset['noi'])} "
                "while maintaining portfolio governance controls and monthly utility review.",
            ],
        )
        add(esg_path, "ESG Snapshot", "Sustainability Office", linked)

        controls_path = target_dir / f"{27 + index:02d}-controls-{index}-{asset['name'].lower().replace(' ', '-')}.pdf"
        _write_pdf(
            controls_path,
            [
                f"Controls and Governance Overview - {asset['name']}",
                f"Reviewed source operating packs, DSCR support, and valuation package for {_money(asset['asset_value'])}.",
                "Control owner confirmed the NOI bridge agrees to the structured quarter state and the "
                "asset-level support schedule used in Winston.",
            ],
        )
        add(controls_path, "Controls + Governance Overview", "Internal Controls", linked)

    # Shared knowledge docs
    data_dictionary = target_dir / "13-data-dictionary-core-tables.docx"
    _write_docx(
        data_dictionary,
        "Core Data Dictionary",
        [
            "Table: re_asset_quarter_state - stores Q1 2026 property NOI, debt, occupancy, and valuation snapshots.",
            "Table: re_fund_quarter_state - stores fund NAV, TVPI, DPI, RVPI, and IRR for Institutional Growth Fund VII.",
            "Table: kb_definition - versioned metric definitions with governance and downstream impact tracking.",
        ],
    )
    add(data_dictionary, "Data Dictionary", "Data Governance", [{"type": "fund", "id": fund["fund_id"]}])

    metric_definitions = target_dir / "14-metric-definitions.pdf"
    _write_pdf(
        metric_definitions,
        [
            "Metric Definitions",
            "NOI = Rental Revenue + Other Income - Vacancy Loss - Operating Expenses.",
            "DSCR = NOI / Debt Service. TVPI = (Distributed Capital + Residual NAV) / Paid-In Capital.",
            "WALT and IRR are tracked alongside valuation to support institutional governance and client reporting.",
        ],
    )
    add(metric_definitions, "Metric Definitions", "Portfolio Analytics", [{"type": "fund", "id": fund["fund_id"]}])

    scenario_playbook = target_dir / "15-scenario-playbook.pdf"
    _write_pdf(
        scenario_playbook,
        [
            "Scenario Playbook",
            "Baseline downside test applies a +75 bps exit cap rate shift.",
            f"The seeded downside case reduces NAV by {_money(abs(fixture['downside']['portfolio_nav_delta']))}, "
            f"changes TVPI by {fixture['downside']['tvpi_delta']:.2f}, and reduces net IRR by "
            f"{fixture['downside']['net_irr_delta'] * 100:.1f} percentage points.",
        ],
    )
    add(scenario_playbook, "Scenario Playbook", "Portfolio Strategy", [{"type": "fund", "id": fund["fund_id"]}])

    lp_extract_1 = target_dir / "32-lp-report-extract-a.csv"
    _write_csv(
        lp_extract_1,
        ["quarter", "fund_name", "portfolio_nav", "tvpi", "net_irr"],
        [[fund["quarter"], fund["name"], str(fund["portfolio_nav"]), f"{fund['tvpi']:.2f}", f"{fund['net_irr']:.3f}"]],
    )
    add(lp_extract_1, "LP Report Extract", "Investor Relations", [{"type": "fund", "id": fund["fund_id"]}])

    lp_extract_2 = target_dir / "33-lp-report-extract-b.csv"
    _write_csv(
        lp_extract_2,
        ["quarter", "total_called", "total_distributed", "dpi", "rvpi"],
        [[fund["quarter"], str(fund["total_called"]), str(fund["total_distributed"]), f"{fund['dpi']:.2f}", f"{fund['rvpi']:.2f}"]],
    )
    add(lp_extract_2, "LP Report Extract", "Investor Relations", [{"type": "fund", "id": fund["fund_id"]}])

    asset_extract = target_dir / "34-asset-kpi-extract.csv"
    _write_csv(
        asset_extract,
        ["asset_name", "quarter", "noi", "asset_value", "debt_balance", "dscr", "walt"],
        [
            [
                asset["name"],
                fund["quarter"],
                str(asset["noi"]),
                str(asset["asset_value"]),
                str(asset["debt_balance"]),
                f"{asset['dscr']:.2f}",
                f"{asset['walt']:.1f}",
            ]
            for asset in assets
        ],
    )
    add(asset_extract, "Asset KPI Extract", "Portfolio Analytics", [{"type": "fund", "id": fund["fund_id"]}])

    return docs
