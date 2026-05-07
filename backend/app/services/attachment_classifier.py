"""Deterministic-first attachment classifier.

Text track: parse stats are computed locally (openpyxl, pypdf, python-docx).
LLM call is optional and limited to the prose summary line, using the model
named by WINSTON_ATTACHMENT_SUMMARY_MODEL. Failure to produce a summary does
not fail the attachment — deterministic stats stand on their own.

Image track: only dimension + mime stats are returned; processing_mode is
always "vision_only". No LLM call at classify time.

Classifier output is persisted on app.document_versions under a metadata
JSONB key "attachment_classification" (if the column exists). The classification
dict is also returned to the caller for immediate use.

IMPORTANT: The classifier may emit suggested write/import actions as strings
only. No write, import, or mutation happens implicitly — those require explicit
user confirmation through audited tool execution.
"""
from __future__ import annotations

import io
import os
import uuid
import logging
from typing import Any

from app.db import get_cursor
from app.repos.supabase_storage_repo import SupabaseStorageRepository

_storage = SupabaseStorageRepository()

logger = logging.getLogger(__name__)

_IMAGE_MIMES = frozenset(
    ["image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif"]
)
_IMAGE_EXTS = frozenset([".png", ".jpg", ".jpeg", ".webp", ".gif"])

_CRM_KEYWORDS = frozenset(
    ["contact", "email", "phone", "company", "stage", "owner", "lead", "opportunity"]
)
_ACCOUNTING_KEYWORDS = frozenset(
    ["debit", "credit", "account", "ledger", "journal", "balance", "revenue", "expense"]
)
_REPE_KEYWORDS = frozenset(
    ["fund", "asset", "noi", "cap_rate", "irr", "tvpi", "investor", "property", "deal"]
)
_DONOR_KEYWORDS = frozenset(
    ["donation", "donor", "pledge", "gift", "campaign", "charity"]
)


def _detect_domain(column_names: list[str]) -> str:
    cols = {c.lower().replace(" ", "_") for c in column_names}
    scores = {
        "crm_export": len(cols & _CRM_KEYWORDS),
        "accounting": len(cols & _ACCOUNTING_KEYWORDS),
        "repe_data": len(cols & _REPE_KEYWORDS),
        "donor_management": len(cols & _DONOR_KEYWORDS),
    }
    best = max(scores, key=lambda k: scores[k])
    return best if scores[best] > 0 else "general"


def _classify_xlsx(content: bytes) -> dict[str, Any]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheet_names = wb.sheetnames
    row_count_total = 0
    all_columns: list[str] = []
    for name in sheet_names:
        ws = wb[name]
        rows = list(ws.iter_rows(values_only=True))
        row_count_total += max(0, len(rows) - 1)  # exclude header row
        if rows:
            headers = [str(c) for c in rows[0] if c is not None]
            all_columns.extend(headers)
    domain = _detect_domain(all_columns)
    return {
        "document_type": "spreadsheet",
        "parser_used": "openpyxl",
        "processing_mode": "text_indexed",
        "parse_stats": {
            "sheet_count": len(sheet_names),
            "sheet_names": sheet_names,
            "row_count_total": row_count_total,
        },
        "detected_domain": domain,
    }


def _classify_pdf(content: bytes) -> dict[str, Any]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    page_count = len(reader.pages)
    text_sample = ""
    char_count = 0
    possible_tables = False
    for page in reader.pages:
        page_text = page.extract_text() or ""
        char_count += len(page_text)
        if len(text_sample) < 2000:
            text_sample += page_text
        # Heuristic: multiple consecutive whitespace-delimited columns suggest tables
        lines = page_text.splitlines()
        multi_col = sum(1 for line in lines if "  " in line and line.strip())
        if multi_col > 3:
            possible_tables = True

    return {
        "document_type": "pdf",
        "parser_used": "pypdf",
        "processing_mode": "text_indexed",
        "parse_stats": {
            "page_count": page_count,
            "char_count": char_count,
            "possible_tables": possible_tables,
            # Never say "contains tables" — only "table-like text detected"
            "table_note": "table-like text detected" if possible_tables else None,
        },
        "_text_sample": text_sample[:2000],
        "detected_domain": "general",
    }


def _classify_docx(content: bytes) -> dict[str, Any]:
    from docx import Document

    doc = Document(io.BytesIO(content))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    char_count = sum(len(p) for p in paragraphs)
    table_count = len(doc.tables)
    text_sample = "\n".join(paragraphs[:20])[:2000]
    return {
        "document_type": "document",
        "parser_used": "python-docx",
        "processing_mode": "text_indexed",
        "parse_stats": {
            "heading_count": len(headings),
            "paragraph_count": len(paragraphs),
            "char_count": char_count,
            "table_count": table_count,
            "first_heading": headings[0] if headings else None,
        },
        "_text_sample": text_sample,
        "detected_domain": "general",
    }


def _classify_csv(content: bytes) -> dict[str, Any]:
    import csv

    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    column_names = rows[0] if rows else []
    domain = _detect_domain(column_names)
    return {
        "document_type": "csv",
        "parser_used": "csv",
        "processing_mode": "text_indexed",
        "parse_stats": {
            "row_count": max(0, len(rows) - 1),
            "column_names": column_names,
        },
        "detected_domain": domain,
    }


def _classify_image(content: bytes, mime_type: str) -> dict[str, Any]:
    try:
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(content))
        width, height = img.size
    except Exception:
        width, height = None, None
    return {
        "document_type": "image",
        "parser_used": "metadata_only",
        "processing_mode": "vision_only",
        "parse_stats": {
            "mime_type": mime_type,
            "byte_size": len(content),
            "width": width,
            "height": height,
        },
        "detected_domain": "image",
    }


def _llm_summary(text_sample: str, document_type: str) -> str | None:
    """One-line prose summary via WINSTON_ATTACHMENT_SUMMARY_MODEL. Never raises."""
    model = os.environ.get("WINSTON_ATTACHMENT_SUMMARY_MODEL", "").strip()
    if not model or not text_sample.strip():
        return None
    try:
        import anthropic
        client = anthropic.Anthropic()
        response = client.messages.create(
            model=model,
            max_tokens=80,
            system="You write single-sentence document summaries. Respond with exactly one sentence, no preamble.",
            messages=[
                {
                    "role": "user",
                    "content": f"Summarize this {document_type} in one sentence:\n\n{text_sample[:1500]}",
                }
            ],
        )
        return response.content[0].text.strip() if response.content else None
    except Exception:
        return None


def _fetch_document_bytes(document_id: uuid.UUID) -> tuple[bytes, str, str]:
    """Fetch latest version bytes from Supabase Storage.

    Returns (bytes, mime_type, original_filename).
    """
    with get_cursor() as cur:
        cur.execute(
            """SELECT dv.bucket, dv.object_key, dv.mime_type, dv.original_filename
               FROM app.document_versions dv
               WHERE dv.document_id = %s
               ORDER BY dv.version_number DESC
               LIMIT 1""",
            (str(document_id),),
        )
        row = cur.fetchone()
    if not row:
        raise LookupError(f"No version found for document_id={document_id}")

    bucket = row["bucket"]
    object_key = row["object_key"]
    mime_type = row.get("mime_type") or ""
    filename = row.get("original_filename") or ""

    # Fetch bytes via signed URL (server-side only — URL never leaves backend)
    signed = _storage.generate_signed_download_url(bucket, object_key)
    import urllib.request
    with urllib.request.urlopen(signed) as resp:
        content = resp.read()

    return content, mime_type, filename


def classify(document_id: uuid.UUID) -> dict[str, Any]:
    """Run deterministic classification on a document.

    Returns the classification dict. Also persists it on the document version
    under metadata key "attachment_classification" if the metadata column exists.
    """
    content, mime_type, filename = _fetch_document_bytes(document_id)
    fname = (filename or "").lower()

    try:
        if mime_type in _IMAGE_MIMES or any(fname.endswith(ext) for ext in _IMAGE_EXTS):
            result = _classify_image(content, mime_type)
        elif mime_type == "application/pdf" or fname.endswith(".pdf"):
            result = _classify_pdf(content)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ) or fname.endswith((".xlsx", ".xls")):
            result = _classify_xlsx(content)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ) or fname.endswith((".docx", ".doc")):
            result = _classify_docx(content)
        elif fname.endswith(".csv"):
            result = _classify_csv(content)
        else:
            result = {
                "document_type": "text",
                "parser_used": "plaintext",
                "processing_mode": "text_indexed",
                "parse_stats": {"char_count": len(content)},
                "_text_sample": content.decode("utf-8", errors="replace")[:2000],
                "detected_domain": "general",
            }
    except Exception as exc:
        logger.warning("attachment_classifier: parse error for %s: %s", document_id, exc)
        result = {
            "document_type": "unknown",
            "parser_used": "failed",
            "processing_mode": "failed",
            "parse_stats": {},
            "detected_domain": "general",
            "error": str(exc)[:200],
        }

    # LLM summary (text track only, never image track, never raises)
    if result.get("processing_mode") == "text_indexed":
        text_sample = result.pop("_text_sample", "")
        result["summary"] = _llm_summary(text_sample, result.get("document_type", "document"))
    else:
        result.pop("_text_sample", None)
        result["summary"] = None

    # Persist on the version record (best-effort; column may not exist yet)
    try:
        import json
        with get_cursor() as cur:
            cur.execute(
                """UPDATE app.document_versions
                   SET metadata = COALESCE(metadata, '{}')::jsonb
                               || %s::jsonb
                   WHERE document_id = %s
                     AND version_number = (
                       SELECT MAX(version_number) FROM app.document_versions
                       WHERE document_id = %s
                     )""",
                (
                    json.dumps({"attachment_classification": result, "processing_mode": result.get("processing_mode")}),
                    str(document_id),
                    str(document_id),
                ),
            )
    except Exception:
        pass  # metadata column may not exist; classification is still returned

    return result
